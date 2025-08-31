# export_report.py
# Generate a Word report with embedded plots & summaries.

import os
import glob
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -------------------------
# Config (paths are relative to current folder)
# -------------------------
MASTER_CSV   = "customer_churn_master.csv"
CLEANED_CSV  = "customer_churn_cleaned.csv"     # optional, if present we'll reference it
PLOTS_DIR    = "eda_plots"
SUM_DIR      = "eda_summaries"
OUT_DOCX     = "Customer_Churn_Data_Preparation_Report.docx"

# Helpful: prioritize a few key plot filenames if present
PREFERRED_PLOTS = [
    "01_churn_distribution.png",
    "02_age_distribution.png",
    "05_correlation_heatmap.png",
    "04_Total_Spend_vs_churn_box.png",
    "04_LoginFrequency_vs_churn_box.png",
    "04_Days_Since_LastLogin_vs_churn_box.png",
]

# -------------------------
# Helpers
# -------------------------
def add_picture_if_exists(doc, image_path, width_inches=6, caption=None):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False

def add_csv_table_preview(doc, csv_path, max_rows=12, max_cols=8, title=None, note=None):
    if not os.path.exists(csv_path):
        return False
    df = pd.read_csv(csv_path)
    if title:
        doc.add_heading(title, level=2)
    # trim size
    df_show = df.iloc[:max_rows, :max_cols].copy()
    # create table
    table = doc.add_table(rows=df_show.shape[0] + 1, cols=df_show.shape[1])
    table.style = "Table Grid"
    # headers
    for j, col in enumerate(df_show.columns):
        table.cell(0, j).text = str(col)
    # rows
    for i in range(df_show.shape[0]):
        for j in range(df_show.shape[1]):
            val = df_show.iat[i, j]
            table.cell(i + 1, j).text = "" if pd.isna(val) else str(val)
    if note:
        doc.add_paragraph(note).italic = True
    return True

def compute_quick_stats(master_csv):
    stats = {}
    if not os.path.exists(master_csv):
        return stats
    df = pd.read_csv(master_csv)
    stats["n_rows"], stats["n_cols"] = df.shape
    if "ChurnStatus" in df.columns:
        churn_counts = df["ChurnStatus"].value_counts(dropna=False).to_dict()
        n1 = churn_counts.get(1, 0)
        n0 = churn_counts.get(0, 0)
        n = n0 + n1
        stats["churn_rate"] = round((n1 / n) * 100, 2) if n > 0 else None
        stats["churn_counts"] = {"retained_0": n0, "churned_1": n1}
    # Simple sanity checks for common features
    for col in ["Total_Spend", "LoginFrequency", "Days_Since_LastLogin", "Txn_Count"]:
        if col in df.columns:
            stats[f"{col}_mean"] = round(df[col].dropna().mean(), 2)
    return stats

# -------------------------
# Build document
# -------------------------
def main():
    doc = Document()
    doc.add_heading("Customer Churn Prediction – Data Preparation Report", level=0)

    # Intro
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This report documents the data preparation work for a customer churn prediction project. "
        "It covers (1) data gathering & integration, (2) exploratory data analysis (EDA), and "
        "(3) data cleaning & preprocessing to produce a machine-learning-ready dataset."
    )

    # Data Gathering
    doc.add_heading("2. Data Gathering", level=1)
    doc.add_paragraph(
        "Source file: Customer_Churn_Data_Large.xlsx (multiple sheets). We used:\n"
        "• Churn_Status – target label (CustomerID, ChurnStatus).\n"
        "• Customer_Demographics – Age, Gender, MaritalStatus, IncomeLevel.\n"
        "• Transaction_History – aggregated per customer into Total_Spend, Avg_Spend, Txn_Count, Distinct_Categories.\n"
        "• Customer_Service – aggregated per customer into Total_Interactions, Resolved_Count, Unresolved_Count, Last_Interaction_Date.\n"
        "• Online_Activity – engineered Days_Since_LastLogin and retained LoginFrequency, ServiceUsage.\n"
    )

    # Quick stats
    stats = compute_quick_stats(MASTER_CSV)
    if stats:
        nrows = stats.get("n_rows")
        ncols = stats.get("n_cols")
        churn_rate = stats.get("churn_rate")
        churn_counts = stats.get("churn_counts", {})
        doc.add_paragraph(
            f"Final master dataset (pre-cleaning): {nrows} rows × {ncols} columns."
        )
        if churn_rate is not None:
            doc.add_paragraph(
                f"Churn rate: {churn_rate}%  "
                f"(retained= {churn_counts.get('retained_0', 0)}, "
                f"churned= {churn_counts.get('churned_1', 0)})."
            )

    # EDA
    doc.add_heading("3. Exploratory Data Analysis (EDA)", level=1)
    doc.add_paragraph(
        "EDA outputs were generated programmatically. Visualizations are in 'eda_plots/' and tabular summaries in "
        "'eda_summaries/'. Below are representative visuals and previews of summary tables."
    )

    # Embed preferred plots (if present), otherwise embed up to 4 PNGs from the folder
    found_any = False
    for fname in PREFERRED_PLOTS:
        path = os.path.join(PLOTS_DIR, fname)
        added = add_picture_if_exists(doc, path, width_inches=6, caption=f"Figure: {fname.replace('_',' ')}")
        found_any = found_any or added

    if not found_any:
        # fallback: any 4 .png in eda_plots
        pngs = sorted(glob.glob(os.path.join(PLOTS_DIR, "*.png")))[:4]
        for p in pngs:
            add_picture_if_exists(doc, p, width_inches=6, caption=f"Figure: {os.path.basename(p)}")

    # Insert tables from EDA summaries (schema_missingness, numeric_describe, categorical_cardinality)
    add_csv_table_preview(
        doc,
        os.path.join(SUM_DIR, "schema_missingness.csv"),
        max_rows=12,
        title="Schema & Missingness (preview)",
        note="(Showing first rows. See full CSV in eda_summaries/.)"
    )
    add_csv_table_preview(
        doc,
        os.path.join(SUM_DIR, "numeric_describe.csv"),
        max_rows=12,
        title="Numeric Describe (preview)",
        note="(Showing first rows. See full CSV in eda_summaries/.)"
    )
    add_csv_table_preview(
        doc,
        os.path.join(SUM_DIR, "categorical_cardinality.csv"),
        max_rows=12,
        title="Categorical Cardinality (preview)",
        note="(Showing first rows. See full CSV in eda_summaries/.)"
    )

    # Cleaning & Preprocessing
    doc.add_heading("4. Data Cleaning & Preprocessing", level=1)
    doc.add_paragraph(
        "We applied a reproducible pipeline:\n"
        "1) Missing values – numeric: median imputation (with missing indicator if >5% missing); "
        "categorical: mode or 'Unknown'.\n"
        "2) Outliers – IQR-based winsorization (clip beyond Q1−1.5×IQR and Q3+1.5×IQR).\n"
        "3) Scaling – z-score standardization for numeric features to prevent scale dominance.\n"
        "4) Encoding – one-hot encoding for low-cardinality categoricals; frequency encoding for high-cardinality.\n"
    )
    if os.path.exists(CLEANED_CSV):
        dfc = pd.read_csv(CLEANED_CSV, nrows=1)
        doc.add_paragraph(
            f"Cleaned dataset file: {CLEANED_CSV} (example columns: {', '.join(list(dfc.columns)[:10])} …)."
        )
    else:
        doc.add_paragraph(
            "Cleaned dataset file: customer_churn_cleaned.csv (generate via clean_preprocess.py)."
        )

    # Deliverables
    doc.add_heading("5. Deliverables", level=1)
    doc.add_paragraph(
        "• customer_churn_master.csv – integrated dataset prior to cleaning.\n"
        "• eda_plots/ – PNG visualizations (churn distribution, distributions, boxplots, heatmap).\n"
        "• eda_summaries/ – CSV tables (schema & missingness, numeric describe, categorical cardinality).\n"
        "• customer_churn_cleaned.csv – cleaned and preprocessed dataset ready for modeling."
    )

    # Next Steps
    doc.add_heading("6. Next Steps", level=1)
    doc.add_paragraph(
        "• Train baseline models (Logistic Regression, Random Forest) and report Accuracy, Precision, Recall, ROC-AUC.\n"
        "• Investigate feature importance (e.g., SHAP) to identify key churn drivers.\n"
        "• Address class imbalance if needed (class weights or SMOTE)."
    )

    # Save
    doc.save(OUT_DOCX)
    print(f"✅ Report saved → {os.path.abspath(OUT_DOCX)}")

if __name__ == "__main__":
    main()
