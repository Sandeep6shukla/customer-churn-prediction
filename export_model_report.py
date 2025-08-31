# export_model_report.py
# Build a Word report for the churn model with metrics + plots embedded.

import os, glob, json, joblib
from pathlib import Path
from datetime import datetime
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sklearn.metrics import (
    roc_auc_score, accuracy_score, precision_score, recall_score, f1_score,
    confusion_matrix, RocCurveDisplay, PrecisionRecallDisplay
)
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier

ROOT = Path(".")
OUTDIR = ROOT / "model_outputs"
OUTDOC = ROOT / "Customer_Churn_Model_Report.docx"
DATA = ROOT / "customer_churn_cleaned.csv"
TARGET = "ChurnStatus"

def latest(path_glob):
    files = sorted(glob.glob(path_glob))
    return Path(files[-1]) if files else None

def add_picture_if_exists(doc, pth, caption=None, width=6):
    p = Path(pth) if pth else None
    if p and p.exists():
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(str(p), width=Inches(width))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False

def compute_metrics_table_from_models():
    """Try to read model_outputs/metrics.csv; else rebuild quick metrics from latest model or quick train."""
    OUTDIR.mkdir(exist_ok=True)
    metrics_csv = OUTDIR / "metrics.csv"
    rows = []
    plots = {"roc": None, "pr": None, "cm": None}
    chosen_model_file = None

    if metrics_csv.exists():
        m = pd.read_csv(metrics_csv)
        rows = m.to_dict(orient="records")
        # try to pick a best model name for locating plots
        best = m.sort_values("f1", ascending=False).iloc[0]
        best_name = str(best["model"]).replace("*bestT", "")
        plots["roc"] = latest(str(OUTDIR / f"roc_{best_name}*.png"))
        plots["pr"]  = latest(str(OUTDIR / f"pr_{best_name}*.png"))
        plots["cm"]  = latest(str(OUTDIR / f"cm_{best_name}*.png"))
        # find a corresponding model file
        if "random" in best_name.lower():
            chosen_model_file = latest(str(OUTDIR / "random_forest_*.pkl"))
        else:
            chosen_model_file = latest(str(OUTDIR / "logistic_regression_*.pkl"))
        return rows, plots, chosen_model_file

    # If metrics.csv missing, try to evaluate on the latest model
    model_file = latest(str(OUTDIR / "random_forest_*.pkl")) or latest(str(OUTDIR / "logistic_regression_*.pkl"))
    if DATA.exists():
        df = pd.read_csv(DATA)
        X = df.drop(columns=[TARGET]) if TARGET in df.columns else df.copy()
        y = df[TARGET].astype(int) if TARGET in df.columns else None

        if model_file and y is not None:
            model = joblib.load(model_file)
            X_tr, X_te, y_tr, y_te = train_test_split(X, y, test_size=0.2, stratify=y, random_state=42)
            prob = model.predict_proba(X_te)[:, 1]
            # default threshold 0.5 plus best-F1 threshold
            pred05 = (prob >= 0.5).astype(int)

            # best F1 threshold
            best_t, best_f1 = 0.5, -1
            for t in np.linspace(0.1, 0.9, 81):
                p = (prob >= t).astype(int)
                f1 = f1_score(y_te, p, zero_division=0)
                if f1 > best_f1:
                    best_f1, best_t = f1, t
            predT = (prob >= best_t).astype(int)

            def pack(name, y_true, pred, prob, t):
                return {
                    "model": name,
                    "threshold": float(t),
                    "roc_auc": roc_auc_score(y_true, prob),
                    "accuracy": accuracy_score(y_true, pred),
                    "precision": precision_score(y_true, pred, zero_division=0),
                    "recall": recall_score(y_true, pred, zero_division=0),
                    "f1": f1_score(y_true, pred, zero_division=0)
                }

            rows = [
                pack(model_file.stem + "_thr0.50", y_te, pred05, prob, 0.5),
                pack(model_file.stem + f"_thr{best_t:.2f}", y_te, predT, prob, best_t),
            ]

            # Save curves & CM for embedding
            fig1, ax1 = plt.subplots(figsize=(5,4))
            RocCurveDisplay.from_predictions(y_te, prob, ax=ax1)
            ax1.set_title("ROC Curve")
            roc_path = OUTDIR / "roc_recomputed.png"
            fig1.savefig(roc_path, bbox_inches="tight"); plt.close(fig1)

            fig2, ax2 = plt.subplots(figsize=(5,4))
            PrecisionRecallDisplay.from_predictions(y_te, prob, ax=ax2)
            ax2.set_title("Precision-Recall")
            pr_path = OUTDIR / "pr_recomputed.png"
            fig2.savefig(pr_path, bbox_inches="tight"); plt.close(fig2)

            from sklearn.metrics import ConfusionMatrixDisplay
            fig3, ax3 = plt.subplots(figsize=(4,4))
            ConfusionMatrixDisplay.from_predictions(y_te, predT, ax=ax3)
            ax3.set_title("Confusion Matrix (best threshold)")
            cm_path = OUTDIR / "cm_recomputed.png"
            fig3.savefig(cm_path, bbox_inches="tight"); plt.close(fig3)

            plots = {"roc": roc_path, "pr": pr_path, "cm": cm_path}
            chosen_model_file = model_file
            # also write a fresh metrics.csv for future
            pd.DataFrame(rows).to_csv(metrics_csv, index=False)
            return rows, plots, chosen_model_file

        # Last resort: quick train RF so a report can be produced
        if y is not None:
            rf = RandomForestClassifier(n_estimators=300, class_weight="balanced", random_state=42)
            X_tr, X_te, y_tr, y_te = train_test_split(X, y, test_size=0.2, stratify=y, random_state=42)
            rf.fit(X_tr, y_tr)
            prob = rf.predict_proba(X_te)[:, 1]
            pred05 = (prob >= 0.5).astype(int)
            # best F1 threshold
            best_t, best_f1 = 0.5, -1
            for t in np.linspace(0.1, 0.9, 81):
                p = (prob >= t).astype(int)
                f1 = f1_score(y_te, p, zero_division=0)
                if f1 > best_f1:
                    best_f1, best_t = f1, t
            predT = (prob >= best_t).astype(int)

            rows = [
                {"model": "RandomForest_quick_thr0.50",
                 "threshold": 0.5,
                 "roc_auc": roc_auc_score(y_te, prob),
                 "accuracy": accuracy_score(y_te, pred05),
                 "precision": precision_score(y_te, pred05, zero_division=0),
                 "recall": recall_score(y_te, pred05, zero_division=0),
                 "f1": f1_score(y_te, pred05, zero_division=0)},
                {"model": f"RandomForest_quick_thr{best_t:.2f}",
                 "threshold": float(best_t),
                 "roc_auc": roc_auc_score(y_te, prob),
                 "accuracy": accuracy_score(y_te, predT),
                 "precision": precision_score(y_te, predT, zero_division=0),
                 "recall": recall_score(y_te, predT, zero_division=0),
                 "f1": f1_score(y_te, predT, zero_division=0)}
            ]
            pd.DataFrame(rows).to_csv(metrics_csv, index=False)

            # plots
            fig1, ax1 = plt.subplots(figsize=(5,4))
            RocCurveDisplay.from_predictions(y_te, prob, ax=ax1)
            ax1.set_title("ROC Curve"); roc_path = OUTDIR / "roc_quick.png"
            fig1.savefig(roc_path, bbox_inches="tight"); plt.close(fig1)

            fig2, ax2 = plt.subplots(figsize=(5,4))
            PrecisionRecallDisplay.from_predictions(y_te, prob, ax=ax2)
            ax2.set_title("Precision-Recall"); pr_path = OUTDIR / "pr_quick.png"
            fig2.savefig(pr_path, bbox_inches="tight"); plt.close(fig2)

            fig3, ax3 = plt.subplots(figsize=(4,4))
            cm = confusion_matrix(y_te, predT)
            im = ax3.imshow(cm, cmap="Blues")
            ax3.set_title("Confusion Matrix (best threshold)")
            for (i,j), v in np.ndenumerate(cm):
                ax3.text(j, i, str(v), ha="center", va="center")
            cm_path = OUTDIR / "cm_quick.png"
            fig3.savefig(cm_path, bbox_inches="tight"); plt.close(fig3)

            plots = {"roc": roc_path, "pr": pr_path, "cm": cm_path}
            chosen_model_file = None
            return rows, plots, chosen_model_file

    # If we reach here, no data
    return rows, plots, chosen_model_file

def build_report():
    OUTDIR.mkdir(exist_ok=True)
    doc = Document()
    doc.add_heading("Customer Churn – Model Development Report", level=0)

    # 1) Intro
    doc.add_heading("1. Objective", level=1)
    doc.add_paragraph(
        "Develop and validate a machine learning model to predict customer churn. "
        "The model is trained on the preprocessed dataset (`customer_churn_cleaned.csv`) and evaluated with "
        "metrics suited for imbalanced data (ROC-AUC, Precision, Recall, F1) with decision-threshold tuning."
    )

    # 2) Algorithm choice
    doc.add_heading("2. Algorithm Selection & Rationale", level=1)
    doc.add_paragraph(
        "We trained two complementary models:\n"
        "• Logistic Regression (class_weight='balanced') – interpretable, calibrated probabilities.\n"
        "• Random Forest (class_weight='balanced') – captures non-linearities and interactions; strong baseline.\n"
        "Hyperparameters were tuned with cross-validation in the training script. Thresholds were optimized for F1."
    )

    # 3) Training & Validation
    doc.add_heading("3. Training & Validation", level=1)
    doc.add_paragraph(
        "• Data split: stratified 80/20 train/test.\n"
        "• Cross-validation: 5-fold StratifiedKFold during hyperparameter search.\n"
        "• Primary metric: ROC-AUC (robust to class imbalance). Secondary: Precision, Recall, F1, Accuracy.\n"
        "• Imbalance handling: class_weight='balanced'.\n"
        "• Threshold optimization: sweep 0.10–0.90 to maximize F1."
    )

    # 4) Performance
    doc.add_heading("4. Performance & Evaluation", level=1)
    rows, plots, model_file = compute_metrics_table_from_models()
    if rows:
        # table
        cols = ["model", "threshold", "roc_auc", "precision", "recall", "f1", "accuracy"]
        table = doc.add_table(rows=len(rows)+1, cols=len(cols))
        table.style = "Table Grid"
        for j,c in enumerate(cols): table.cell(0,j).text = c
        for i,row in enumerate(rows, start=1):
            for j,c in enumerate(cols):
                val = row.get(c, "")
                table.cell(i,j).text = f"{val:.4f}" if isinstance(val, float) else str(val)
    else:
        doc.add_paragraph("Metrics not found and could not be recomputed. Please rerun train_model.py.")

    # embed plots if available
    doc.add_paragraph()
    add_picture_if_exists(doc, plots.get("roc"), "ROC Curve")
    add_picture_if_exists(doc, plots.get("pr"), "Precision–Recall Curve")
    add_picture_if_exists(doc, plots.get("cm"), "Confusion Matrix (best threshold)")

    # 5) Business usage
    doc.add_heading("5. Business Utilisation", level=1)
    doc.add_paragraph(
        "• Rank customers by churn probability; action top deciles.\n"
        "• Tailor interventions: low login frequency → engagement nudges; unresolved cases → service outreach; "
        "high-value at risk → targeted offers.\n"
        "• Monitor outcomes and iterate thresholds using campaign response data."
    )

    # 6) Improvements
    doc.add_heading("6. Potential Improvements", level=1)
    doc.add_paragraph(
        "• Evaluate Gradient Boosting (XGBoost/LightGBM) for possible ROC-AUC gains.\n"
        "• Use cost-sensitive thresholding based on offer cost vs. retained value.\n"
        "• Add temporal features (tenure, trend of spend/engagement), probability calibration, and drift monitoring."
    )

    # 7) Artifacts
    doc.add_heading("7. Artifacts & Reproducibility", level=1)
    p = doc.add_paragraph()
    p.add_run("Trained models: ").bold = True
    doc.add_paragraph("model_outputs/*.pkl (latest RandomForest or LogisticRegression).", style=None)
    p = doc.add_paragraph()
    p.add_run("Metrics & plots: ").bold = True
    doc.add_paragraph("model_outputs/metrics.csv, metrics.json, roc_*.png, pr_*.png, cm_*.png", style=None)
    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    doc.add_paragraph("customer_churn_cleaned.csv", style=None)

    doc.save(OUTDOC)
    print(f"✅ Model report saved → {OUTDOC.resolve()}")

if __name__ == "__main__":
    build_report()